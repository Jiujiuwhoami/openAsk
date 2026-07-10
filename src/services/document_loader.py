"""文档加载器工厂：根据文件扩展名返回对应的加载器。"""

import os
from typing import Optional, Type

from pypdf import PdfReader

from src.domain.exceptions import KnowledgeBaseError
from src.utils.logger import get_logger

logger = get_logger(__name__)


class DocumentLoaderFactory:
    """文档加载器工厂类。

    使用工厂模式，根据文件扩展名自动选择合适的加载器：
    - .md → MarkdownLoader
    - .txt → TextLoader
    - .pdf → PyPDFLoader
    - .docx → DocxLoader
    - .html → HtmlLoader
    - .png/.jpg/.jpeg/.gif/.webp → ImageLoader

    Examples:
        >>> factory = DocumentLoaderFactory()
        >>> loader = factory.get_loader("data/documents/faq/001.md")
        >>> docs = loader.load()
    """

    _LOADER_MAP = {
        ".md": "MarkdownLoader",
        ".txt": "TextLoader",
        ".docx": "DocxLoader",
        ".html": "HtmlLoader",
        ".pdf": "PyPDFLoader",
        ".png": "ImageLoader",
        ".jpg": "ImageLoader",
        ".jpeg": "ImageLoader",
        ".gif": "ImageLoader",
        ".webp": "ImageLoader",
    }

    @classmethod
    def get_loader(cls, file_path: str):
        """根据文件扩展名获取对应的加载器。

        Args:
            file_path: 文档文件路径

        Returns:
            对应的 DocumentLoader 实例

        Raises:
            KnowledgeBaseError: 文件不存在或扩展名不支持
        """
        if not os.path.exists(file_path):
            raise KnowledgeBaseError(f"文件不存在: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()

        if ext in cls._LOADER_MAP:
            loader_name = cls._LOADER_MAP[ext]
            loader_class = globals()[loader_name]
            logger.debug(f"使用 {loader_name} 加载文件: {file_path}")
            return loader_class(file_path)

        raise KnowledgeBaseError(f"不支持的文件格式: {ext}")

    @classmethod
    def supports(cls, file_path: str) -> bool:
        """检查文件扩展名是否支持。"""
        ext = os.path.splitext(file_path)[1].lower()
        return ext in cls._LOADER_MAP


class MarkdownLoader:
    """Markdown 文件加载器。"""

    def __init__(self, file_path: str):
        self._file_path = file_path

    def load(self):
        """加载 Markdown 文件，返回 Document 列表。"""
        from langchain_core.documents import Document

        try:
            with open(self._file_path, "r", encoding="utf-8") as f:
                content = f.read()

            docs = [
                Document(
                    page_content=content,
                    metadata={
                        "source": self._file_path,
                        "title": os.path.splitext(os.path.basename(self._file_path))[0],
                    },
                )
            ]
            logger.debug(f"Markdown 文件加载完成: {self._file_path}")
            return docs
        except Exception as e:
            raise KnowledgeBaseError(f"Markdown 文件加载失败: {e}")


class TextLoader:
    """文本文件加载器。"""

    def __init__(self, file_path: str):
        self._file_path = file_path

    def load(self):
        """加载文本文件，返回 Document 列表。"""
        from langchain_core.documents import Document

        try:
            with open(self._file_path, "r", encoding="utf-8") as f:
                content = f.read()

            docs = [
                Document(
                    page_content=content,
                    metadata={
                        "source": self._file_path,
                        "title": os.path.splitext(os.path.basename(self._file_path))[0],
                    },
                )
            ]
            logger.debug(f"文本文件加载完成: {self._file_path}")
            return docs
        except Exception as e:
            raise KnowledgeBaseError(f"文本文件加载失败: {e}")


class DocxLoader:
    """Word 文档加载器。"""

    def __init__(self, file_path: str):
        self._file_path = file_path

    def load(self):
        """加载 Word 文档，返回 Document 列表。"""
        from langchain_core.documents import Document

        try:
            import docx

            doc = docx.Document(self._file_path)
            content = "\n".join([para.text for para in doc.paragraphs])

            docs = [
                Document(
                    page_content=content,
                    metadata={
                        "source": self._file_path,
                        "title": os.path.splitext(os.path.basename(self._file_path))[0],
                    },
                )
            ]
            logger.debug(f"Word 文档加载完成: {self._file_path}")
            return docs
        except ImportError:
            raise KnowledgeBaseError("需要安装 python-docx 库：pip install python-docx")
        except Exception as e:
            raise KnowledgeBaseError(f"Word 文档加载失败: {e}")


class HtmlLoader:
    """HTML 文件加载器。"""

    def __init__(self, file_path: str):
        self._file_path = file_path

    def load(self):
        """加载 HTML 文件，提取文本内容。"""
        from langchain_core.documents import Document

        try:
            from bs4 import BeautifulSoup

            with open(self._file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "lxml")

            content = soup.get_text(strip=True)

            docs = [
                Document(
                    page_content=content,
                    metadata={
                        "source": self._file_path,
                        "title": soup.title.string if soup.title else "",
                    },
                )
            ]
            logger.debug(f"HTML 文件加载完成: {self._file_path}")
            return docs
        except ImportError:
            raise KnowledgeBaseError("需要安装 beautifulsoup4 和 lxml 库")
        except Exception as e:
            raise KnowledgeBaseError(f"HTML 文件加载失败: {e}")


class PyPDFLoader:
    """PDF 文件加载器封装。

    使用 pypdf 库读取 PDF 文件内容，将每页转换为一个 Document。
    """

    def __init__(self, file_path: str):
        self._file_path = file_path

    def load(self):
        """加载 PDF 文件，返回 Document 列表。"""
        from langchain_core.documents import Document

        docs = []
        try:
            reader = PdfReader(self._file_path)
            for page_num, page in enumerate(reader.pages, 1):
                content = page.extract_text() or ""
                if content.strip():
                    docs.append(
                        Document(
                            page_content=content,
                            metadata={
                                "source": self._file_path,
                                "page": page_num,
                                "total_pages": len(reader.pages),
                            },
                        )
                    )
            logger.debug(f"PDF 文件加载完成: {self._file_path}（共 {len(docs)} 页）")
            return docs
        except Exception as e:
            raise KnowledgeBaseError(f"PDF 文件加载失败: {e}")


class ImageLoader:
    """图片加载器：使用多模态服务描述图片内容。

    使用云端多模态 API（如 OpenAI GPT-4V）识别图片中的内容，
    将识别结果作为文档内容返回，以便后续向量化和检索。

    支持两种初始化方式：
    1. 自动创建：只传入 file_path，内部自动获取多模态服务
    2. 依赖注入：传入 file_path 和 multimodal_service，便于测试和复用

    Examples:
        >>> loader = ImageLoader("screenshot.png")
        >>> docs = loader.load()
        >>> print(docs[0].page_content)  # 图片描述文本
    """

    def __init__(
        self,
        file_path: str,
        multimodal_service=None,
    ):
        self._file_path = file_path
        self._multimodal_service = multimodal_service

    def load(self):
        """加载图片并识别内容，返回 Document 列表。"""
        from langchain_core.documents import Document

        try:
            service = self._get_service()
            description = service.describe_image(self._file_path)

            docs = [
                Document(
                    page_content=description,
                    metadata={
                        "source": self._file_path,
                        "type": "image",
                    },
                )
            ]
            logger.debug(f"图片文件加载完成: {self._file_path}（识别结果 {len(description)} 字符）")
            return docs
        except Exception as e:
            raise KnowledgeBaseError(f"图片识别失败: {e}")

    def _get_service(self):
        """获取多模态服务实例（延迟加载）。"""
        if self._multimodal_service is not None:
            return self._multimodal_service

        from src.infrastructure.multimodal_service import MultiModalServiceFactory

        return MultiModalServiceFactory.get_service()